"""End-to-end upload pipeline test: Go -> Python -> ChromaDB."""
import os
import sys
import time

sys.stdout.reconfigure(encoding='utf-8')
sys.stderr.reconfigure(encoding='utf-8')

from dotenv import load_dotenv
load_dotenv()
os.environ["HF_ENDPOINT"] = os.getenv("HF_ENDPOINT", "https://hf-mirror.com")

import chromadb
from chromadb.config import Settings
from docx import Document

PYTHON_HOST = os.getenv("PYTHON_HOST", "localhost")
PYTHON_PORT = os.getenv("PYTHON_PORT", "5000")
GO_PORT = os.getenv("SERVER_PORT", "8080")

BASE_URL = f"http://localhost:{GO_PORT}/api/v1"

def create_test_docx():
    """Create a test DOCX with reimbursement-related content."""
    doc = Document()
    para = doc.add_paragraph()
    run = para.add_run()
    run.text = (
        "报销申请单：\n"
        "员工姓名：张三\n"
        "部门：销售部\n"
        "报销类型：差旅费\n"
        "报销金额：1500元\n"
        "发票编号：FP20240101001\n"
        "审批状态：待审批\n\n"
        "明细：\n"
        "1. 机票费用：800元\n"
        "2. 酒店住宿：500元\n"
        "3. 交通补助：200元"
    )
    path = "/tmp/test_reimbursement_baoxiao.docx"
    doc.save(path)
    return path


def check_chroma_count():
    """Check how many chunks are in ChromaDB."""
    client = chromadb.PersistentClient(
        path="./vector_db",
        settings=Settings(anonymized_telemetry=False)
    )
    try:
        collection = client.get_collection(name="lumina_docs")
        return collection.count()
    except:
        return 0


def main():
    print("=" * 60)
    print("Upload Pipeline Test")
    print("=" * 60)

    before_count = check_chroma_count()
    print(f"\n[1] ChromaDB before upload: {before_count} chunks")

    # Create test file
    docx_path = create_test_docx()
    print(f"[2] Created test DOCX: {docx_path}")

    # Upload via Go backend
    print(f"[3] Uploading to Go backend at {BASE_URL}...")
    import requests

    with open(docx_path, 'rb') as f:
        resp = requests.post(
            f"{BASE_URL}/docs/upload",
            files={'file': ('报销测试.docx', f, 'application/vnd.openxmlformats-officedocument.wordprocessingml.document')}
        )

    print(f"    Response status: {resp.status_code}")
    print(f"    Response body: {resp.text[:300]}")

    if resp.status_code not in (200, 201):
        print("[ERROR] Upload failed!")
        return

    # Wait for async ingestion
    print("[4] Waiting for async ingestion (5s)...")
    time.sleep(5)

    after_count = check_chroma_count()
    print(f"\n[5] ChromaDB after upload: {after_count} chunks")

    if after_count > before_count:
        print(f"\n[PASS] Ingestion SUCCESS! Added {after_count - before_count} chunks")
    else:
        print(f"\n[FAIL] No new chunks added. DB still has {after_count} chunks")


if __name__ == "__main__":
    main()